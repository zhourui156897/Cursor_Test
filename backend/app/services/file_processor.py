"""Multimodal file processor: extract text from any file format → pure Markdown.

Supports: text, markdown, docx, xlsx, pdf, images (via OpenAI Vision),
audio (via OpenAI Whisper), video (extract audio with ffmpeg → Whisper).
Falls back gracefully when deps are missing. Video requires ffmpeg on PATH.
"""

from __future__ import annotations

import base64
import logging
import mimetypes
from pathlib import Path

logger = logging.getLogger(__name__)

# Register extra MIME types that Python doesn't know
mimetypes.add_type("application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx")
mimetypes.add_type("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", ".xlsx")
mimetypes.add_type("application/vnd.openxmlformats-officedocument.presentationml.presentation", ".pptx")


async def extract_text(file_path: str | Path, content_type: str | None = None) -> str:
    """Extract text from any file and return as Markdown."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    ext = path.suffix.lower()

    # Always re-detect from extension — browser content_type is unreliable
    guessed, _ = mimetypes.guess_type(str(path))
    if content_type in (None, "application/octet-stream", ""):
        content_type = guessed or "application/octet-stream"

    # Text-based files (check extension first since MIME detection is unreliable)
    TEXT_EXTS = {".txt", ".md", ".markdown", ".csv", ".html", ".htm", ".json", ".xml", ".yaml", ".yml", ".log", ".ini", ".cfg", ".conf", ".toml", ".rst", ".tex"}
    if ext in TEXT_EXTS or content_type in ("text/plain", "text/markdown", "text/csv", "text/html", "application/json"):
        return _read_text(path)

    # Word documents
    if ext in (".docx",) or "wordprocessingml" in content_type:
        return _extract_docx(path)

    # Excel
    if ext in (".xlsx",) or "spreadsheetml" in content_type:
        return _extract_xlsx(path)

    # PDF
    if ext == ".pdf" or content_type == "application/pdf":
        return _extract_pdf(path)

    # Images → OpenAI Vision
    if content_type.startswith("image/") or ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"):
        return await _extract_image_vision(path)

    # Audio → OpenAI Whisper
    if content_type.startswith("audio/") or ext in (".mp3", ".wav", ".m4a", ".ogg", ".flac", ".webm"):
        return await _extract_audio_whisper(path)

    # Video — 抽音轨 → Whisper 转文字（需本机安装 ffmpeg）
    if content_type.startswith("video/") or ext in (".mp4", ".mov", ".avi", ".mkv", ".webm"):
        return await _extract_video_via_whisper(path)

    return f"[文件: {path.name}] (不支持的格式: {content_type})"


def _read_text(path: Path) -> str:
    for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_docx(path: Path) -> str:
    """Extract text from Word .docx → Markdown."""
    try:
        from docx import Document
    except ImportError:
        return f"[Word文件: {path.name}] (需要安装: pip install python-docx)"

    doc = Document(str(path))
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        style = para.style.name if para.style else ""
        if "Heading 1" in style:
            lines.append(f"# {text}")
        elif "Heading 2" in style:
            lines.append(f"## {text}")
        elif "Heading 3" in style:
            lines.append(f"### {text}")
        elif "List" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, header_sep)
            lines.append("\n" + "\n".join(rows))

    return "\n".join(lines).strip()


def _extract_xlsx(path: Path) -> str:
    """Extract tables from Excel .xlsx → Markdown tables."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return f"[Excel文件: {path.name}] (需要安装: pip install openpyxl)"

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts = []

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        parts.append(f"## {sheet}\n")
        for i, row in enumerate(rows[:200]):
            cells = [str(c) if c is not None else "" for c in row]
            parts.append("| " + " | ".join(cells) + " |")
            if i == 0:
                parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
        parts.append("")

    wb.close()
    return "\n".join(parts).strip() or f"[Excel文件: {path.name}] (空文件)"


def _extract_pdf(path: Path) -> str:
    """Extract text from PDF → Markdown."""
    try:
        import fitz
    except ImportError:
        return f"[PDF文件: {path.name}] (需要安装: pip install pymupdf)"

    doc = fitz.open(str(path))
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text().strip()
        if text:
            pages.append(f"<!-- Page {i+1} -->\n{text}")
    doc.close()

    if pages:
        return "\n\n---\n\n".join(pages)
    return f"[PDF文件: {path.name}] (无法提取文本，可能是扫描件，建议截图后上传图片)"


async def _extract_image_vision(path: Path) -> str:
    """Use OpenAI Vision API to OCR/describe an image."""
    from app.config import get_settings, get_user_config

    settings = get_settings()
    cfg = get_user_config()
    api_key = cfg.get("llm", {}).get("api_key", "")
    api_url = settings.llm_api_url

    if not api_key:
        return f"[图片: {path.name}] (需要配置 LLM API Key 才能进行图片识别)"

    data = base64.b64encode(path.read_bytes()).decode("utf-8")
    ext = path.suffix.lower().lstrip(".")
    mime_map = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "gif": "gif", "webp": "webp", "bmp": "bmp"}
    mime = f"image/{mime_map.get(ext, ext)}"

    import httpx
    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.post(
            f"{api_url.rstrip('/')}/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": settings.llm_model,
                "messages": [
                    {"role": "system", "content": "你是OCR和图片分析助手。请识别图片中的所有文字并保持排版。如果没有文字，详细描述图片内容。输出纯 Markdown 格式。"},
                    {"role": "user", "content": [
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{data}"}},
                        {"type": "text", "text": "请识别图片中的文字内容，转为 Markdown 格式输出。"},
                    ]},
                ],
                "max_tokens": 4096,
                "temperature": 0.1,
            },
        )
        resp.raise_for_status()
        result = resp.json()
        return result["choices"][0]["message"]["content"]


async def _extract_video_via_whisper(path: Path) -> str:
    """Extract audio from video with ffmpeg, then transcribe with Whisper. Requires ffmpeg on PATH."""
    import shutil
    import tempfile

    if not shutil.which("ffmpeg"):
        return f"[视频: {path.name}] (视频转文字需要本机安装 ffmpeg，请先安装: brew install ffmpeg)"

    with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        import asyncio

        proc = await asyncio.create_subprocess_exec(
            "ffmpeg", "-y", "-i", str(path), "-vn", "-acodec", "libmp3lame", "-q:a", "2", str(tmp_path),
            stdout=asyncio.subprocess.DEVNULL,
            stderr=asyncio.subprocess.PIPE,
        )
        _, stderr = await proc.communicate()
        if proc.returncode != 0 or not tmp_path.exists():
            logger.warning("ffmpeg extract audio failed: %s", stderr.decode()[:500] if stderr else "")
            return f"[视频: {path.name}] (无法从视频提取音轨，请检查 ffmpeg 或文件格式)"

        text = await _extract_audio_whisper(tmp_path)
        return f"**来自视频转写** ({path.name})\n\n{text}" if text else f"[视频: {path.name}] (转写无内容)"
    finally:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except OSError:
                pass


async def _extract_audio_whisper(path: Path) -> str:
    """Use OpenAI Whisper API to transcribe audio."""
    from app.config import get_settings, get_user_config

    settings = get_settings()
    cfg = get_user_config()
    api_key = cfg.get("llm", {}).get("api_key", "")

    if not api_key:
        return f"[音频: {path.name}] (需要配置 OpenAI API Key 才能进行语音转文字)"

    # Only OpenAI has Whisper endpoint
    if "openai.com" not in settings.llm_api_url:
        return f"[音频: {path.name}] (语音转文字需要 OpenAI API，当前 API 不支持)"

    import httpx
    async with httpx.AsyncClient(timeout=120.0) as client:
        with open(path, "rb") as f:
            resp = await client.post(
                "https://api.openai.com/v1/audio/transcriptions",
                headers={"Authorization": f"Bearer {api_key}"},
                files={"file": (path.name, f, "audio/mpeg")},
                data={"model": "whisper-1", "response_format": "text", "language": "zh"},
            )
        resp.raise_for_status()
        return resp.text.strip()


def detect_content_type(filename: str) -> str:
    """Detect content type category from filename."""
    ext = Path(filename).suffix.lower()
    ct, _ = mimetypes.guess_type(filename)
    ct = ct or "application/octet-stream"

    if ext in (".docx",):
        return "text/markdown"
    if ext in (".xlsx",):
        return "text/markdown"
    if ct.startswith("text/") or ct in ("application/json", "application/xml"):
        return "text"
    if ct == "application/pdf":
        return "pdf"
    if ct.startswith("image/"):
        return "image"
    if ct.startswith("audio/"):
        return "audio"
    if ct.startswith("video/"):
        return "video"
    return "file"
